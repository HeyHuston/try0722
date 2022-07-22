from docx import Document
from docx.shared import Pt#设置字体
from docx.oxml.ns import qn#设置中文字体
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import pandas as pd

n=0;colorindex=1;hhh=""
document = Document("tyr.docx")
document1=Document();document2=Document();document3=Document();document4=Document();document5=Document();document6=Document();document7=Document();document8=Document();document9=Document();document10=Document()
doclist=[document1,document2,document3,document4,document5,document6,document7,document8,document9,document10] 
color=[RGBColor(50,200,150),RGBColor(50,150,200),RGBColor(200,150,50),RGBColor(150,200,50),RGBColor(200,50,150),RGBColor(150,50,200),RGBColor(150,50,255),RGBColor(255,50,150),RGBColor(50,255,150),RGBColor(50,150,255),]

Qun=pd.DataFrame(columns=['群名','doc'])
PL=pd.DataFrame(columns=['QQ号','颜色']) 
PL.loc['Aldebaran']=['2313677449',color[0]]
doc_dict={};pl_dict={}


def gaiming(pl):
    if (pl=="冬客"):
        pl="严景灏"
    if (pl=="懒得想名字了"):
        pl="姚通海"
    if (pl=="白河寒秋"):
        pl="余一唯"
    if (pl=="Nof"):
        pl="苗梧"
    if (pl=="余愚余娱"):
        pl="纪舒云"
    if (pl=="阿遥"):
        pl="幽灵"
    if (pl=="远飞的足球<LTR>"):
        pl="幽灵2号"
    if (pl=="Tina Tang"):
        pl="Keeper"
    return pl

def chu(text):
    m=text.find('(')
    n=text.find(')')    
    num=text[m+1:n]
    qun=text[:m]
    return num,qun
def chu2(text):
    m=text.find('[')
    n=text.find(']')    
    text=text[m+1:n]
    return text
 
biaozhici1=["Friend(2313677449)","Event:","DiceDriver","Connection lost","Logging in","successful",'Reconnected','Aldebaran、懒得想名字了'] 
for no in range(0,len(document.paragraphs)):
    text=document.paragraphs[no].text
    if any (s in text for s in biaozhici1):
        document.paragraphs[no].clear()
document.save("gai1.docx")
document=Document('gai1.docx')
for para in document.paragraphs:
    num=''
    text=para.text            
    flag=0;qun1=0;m=0;qun2=0;pl1=0;pl2=0;time1=0
    
    if (text==''):
        continue
    else:
        if any (s in text for s in ["Group"]):
            time1=text.find(' ')            
            pl2=text.find('<')
            time=text[time1:time1+8];pc=text[pl2+2:];pl="Aldebaran";num=text[pl2-12:pl2-3];num=chu(text)[0]
            he=time+" "+pl+"："+pc
        elif any (s in text for s in ["Friend(335904632"]):
            time1=text.find(' ')            
            pl2=text.find('<')
            time=text[time1:time1+8];pc=text[pl2+2:];pl="Aldebaran";num=text[pl2-12:pl2-3];num=chu(pc)[0]
            he=time+" "+pl+"：Keeper"+pc
        else:
            flag_qun=0;m=0;flag=0;qun1=0;qun2=0;pl1=0;pl2=0;time1=0
           
            time=text[time1:time1+8];qun=text[qun1:qun2-1];pl=text[qun2+1:pl1-1];pc=text[pl2+1:]
            num=chu(qun)[0];qun=chu(qun)[1]
            plnum=chu(pl)[0];pl=chu(pl)[1]
            pl=gaiming(pl)
            pc=chu2(pc)
            if any(num in nums for nums in Qun.index):
                hhh=""
            else:
                Qun.loc[num]=[qun,doclist[n]]
                n=n+1
            if any(pl in pls for pls in PL.index):
                hhh=""
            else:
                PL.loc[pl]=[plnum,color[colorindex]]
                colorindex=colorindex+1
                print(PL)
            he= he=time+" "+pl+"："+pc
            
        if (pc==''):
            continue
        else:
            run = Qun.at[num,'doc'].add_paragraph().add_run(he)
            print(text)
            run.font.color.rgb =PL.at[pl,'颜色']
print(Qun)
print(PL)         
for num in Qun.index:
    add=Qun.at[num,'群名']+".docx" 
    Qun.at[num,'doc'].save(add)
    


            
        
            
