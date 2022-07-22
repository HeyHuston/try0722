from docx import Document
import pandas as pd
from openpyxl import load_workbook

def gaiming(pl):
    if (pl=="冬客"):
        pl="严景灏"
    if (pl=="懒得想名字了"):
        pl="姚通海"
    if (pl=="白河寒秋"):
        pl="余一唯"
    if (pl=="Nof"):
        pl="苗梧"
    if (pl=="余愚余娱"):
        pl="纪舒云"
    if (pl=="阿遥"):
        pl="幽灵"
    if (pl=="远飞的足球<LTR>"):
        pl="幽灵2号"
    if (pl=="Tina Tang"):
        pl="Keeper"
    if (pl=="阿尔瓦"):
        pl="Keeper"
    return pl

i=0;m=0;n=0;p=0;q=0;num=0
DF1=pd.DataFrame();DF2=pd.DataFrame();DF3=pd.DataFrame();DF4=pd.DataFrame();DF5=pd.DataFrame();DF6=pd.DataFrame();DF7=pd.DataFrame();DF8=pd.DataFrame();DF9=pd.DataFrame();DF10=pd.DataFrame()
DFlist=[DF1,DF2,DF3,DF4,DF5,DF5,DF7,DF8,DF9,DF10]
for i in range(0,len(DFlist)):
    DFlist[i]=pd.DataFrame(columns=['技能值','尝试次数','成功次数','成功率','大成功次数','大失败次数'])
pclist=[]
pcdic={}
document=Document("合.docx")
for no in range(0,len(document.paragraphs)):
    pc='';jn='';jnz=''
    text=document.paragraphs[no].text
    Flag1='Aldebaran'in text
    Flag2='检定'in text
    Flag3='默认检定房规已设置'in text
    if (Flag1==True)and(Flag2==True)and(Flag3==False):
        i=text.find('：')
        m=text.find('进行')
        n=text.find('检定')
        p=text.find('/')
        q=text.find(' ',20)
        cg='成功' in text
        dcg='大成功' in text
        dsb='大失败' in text
        pc=text[i+1:m];jn=text[m+2:n];jnz=text[p+1:q]
        pc=gaiming(pc)
        if any(pc in pcs for pcs in pclist):
            over=''
        else:
            pclist.append(pc)
            pcdic[pc]=DFlist[num]
            print(pc)
            num=num+1
        DF=pcdic[pc]
        if any(jn == jns for jns in DF.index):
            DF.at[jn,'尝试次数']=DF.at[jn,'尝试次数']+1
            if (cg==True):
                DF.at[jn,'成功次数']=DF.at[jn,'成功次数']+1     
            if (dcg==True):
                DF.at[jn,'大成功次数']=DF.at[jn,'大成功次数']+1
            if (dsb==True):
                DF.at[jn,'大失败次数']=DF.at[jn,'大失败次数']+1
            DF.at[jn,'成功率']=DF.at[jn,'成功次数']/DF.at[jn,'尝试次数']
        else:
            DF.loc[jn]=[jnz,1,0,0,0,0]
            if (cg==True):
                DF.at[jn,'成功次数']=DF.at[jn,'成功次数']+1     
            if (dcg==True):
                DF.at[jn,'大成功次数']=DF.at[jn,'大成功次数']+1
            if (dsb==True):
                DF.at[jn,'大失败次数']=DF.at[jn,'大失败次数']+1
            DF.at[jn,'成功率']=DF.at[jn,'成功次数']/DF.at[jn,'尝试次数']

pcdic['Keeper'].to_excel("表.xlsx",sheet_name='1')       
for pc in pclist:
    DF=pcdic[pc]
    #add=str(pc)+"表.xlsx"
    #DF.to_excel("表.xlsx",sheet_name=pc)
    book = load_workbook('表.xlsx')
    writer = pd.ExcelWriter("表.xlsx",engine='openpyxl')
    writer.book = book
    DF.to_excel(writer,pc)
    writer.save()
        
